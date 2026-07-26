"""
文档加载器：处理 PDF、Word、Excel 三类文档
"""
import json
import re
from pathlib import Path
from typing import List, Dict, Optional

import pdfplumber
from docx import Document
import openpyxl
import jieba


class DocumentChunk:
    """文档块"""
    def __init__(self, content: str, metadata: Dict):
        self.content = content
        self.metadata = metadata

    def __repr__(self):
        return f"Chunk(source={self.metadata.get('source')}, len={len(self.content)})"


class DocumentLoader:
    """统一文档加载器"""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    # ===== PDF =====
    def load_pdf(self, file_path: Path) -> List[DocumentChunk]:
        """加载PDF，保留表格结构"""
        chunks = []
        source = file_path.name

        with pdfplumber.open(str(file_path)) as pdf:
            full_text = []
            for page_num, page in enumerate(pdf.pages, 1):
                # 提取文本
                text = page.extract_text()
                if text:
                    full_text.append(text)

                # 提取表格
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables):
                    table_text = self._table_to_text(table, source, page_num, table_idx)
                    if table_text:
                        full_text.append(table_text)

            content = "\n".join(full_text)
            chunks = self._split_text(content, source, "pdf")

        return chunks

    # ===== Word =====
    def load_docx(self, file_path: Path) -> List[DocumentChunk]:
        """加载Word文档"""
        doc = Document(str(file_path))
        source = file_path.name

        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                # 保留标题层级信息
                if para.style.name.startswith("Heading"):
                    full_text.append(f"## {para.text.strip()}")
                else:
                    full_text.append(para.text.strip())

        # 提取表格
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = f"[表格{table_idx+1}]\n" + "\n".join(rows)
            full_text.append(table_text)

        content = "\n".join(full_text)
        return self._split_text(content, source, "docx")

    # ===== Excel =====
    def load_excel(self, file_path: Path) -> List[DocumentChunk]:
        """加载Excel，每行作为一个结构化记录"""
        wb = openpyxl.load_workbook(str(file_path), data_only=True)
        source = file_path.name
        all_records = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue

            headers = [str(h) if h else "" for h in rows[0]]
            records = []

            for row in rows[1:]:
                if any(cell is not None for cell in row):
                    record = {}
                    for i, value in enumerate(row):
                        if i < len(headers) and headers[i]:
                            record[headers[i]] = str(value) if value is not None else ""
                    if record:
                        records.append(record)
                        # 每行也生成文本表示
                        text = f"[{sheet_name}] " + "，".join(
                            f"{k}: {v}" for k, v in record.items() if v
                        )
                        all_records.append(DocumentChunk(
                            content=text,
                            metadata={
                                "source": source,
                                "sheet": sheet_name,
                                "type": "structured_data",
                                "record": json.dumps(record, ensure_ascii=False)
                            }
                        ))

        return all_records

    # ===== 通用加载入口 =====
    def load_file(self, file_path: Path) -> List[DocumentChunk]:
        """根据文件类型自动选择加载器"""
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self.load_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            return self.load_docx(file_path)
        elif suffix in (".xlsx", ".xls"):
            return self.load_excel(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {suffix}")

    def load_directory(self, dir_path: Path) -> List[DocumentChunk]:
        """批量加载目录下所有文档"""
        all_chunks = []
        for file_path in sorted(dir_path.iterdir()):
            if file_path.suffix.lower() in (".pdf", ".docx", ".doc", ".xlsx", ".xls"):
                try:
                    chunks = self.load_file(file_path)
                    all_chunks.extend(chunks)
                    print(f"  ✅ {file_path.name} → {len(chunks)} chunks")
                except Exception as e:
                    print(f"  ❌ {file_path.name}: {e}")
        return all_chunks

    # ===== 内部分块 =====
    def _split_text(self, text: str, source: str, doc_type: str) -> List[DocumentChunk]:
        """按chunk_size滑动分块"""
        chunks = []
        paragraphs = text.split("\n")
        current = ""
        current_start = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if len(current) + len(para) <= self.chunk_size:
                current += para + "\n"
            else:
                if current.strip():
                    chunks.append(DocumentChunk(
                        content=current.strip(),
                        metadata={"source": source, "type": doc_type}
                    ))
                # 重叠：保留最后一部分
                overlap_text = current[-self.chunk_overlap:] if len(current) > self.chunk_overlap else ""
                current = overlap_text + para + "\n"

        if current.strip():
            chunks.append(DocumentChunk(
                content=current.strip(),
                metadata={"source": source, "type": doc_type}
            ))

        return chunks

    # ===== 表格转文本 =====
    def _table_to_text(self, table, source: str, page: int, table_idx: int) -> str:
        """PDF表格转可读文本"""
        if not table or len(table) < 2:
            return ""

        lines = []
        headers = [str(h) if h else "" for h in table[0]]

        for row in table[1:]:
            if any(cell is not None for cell in row):
                parts = []
                for i, cell in enumerate(row):
                    if i < len(headers) and cell is not None:
                        parts.append(f"{headers[i]}: {cell}")
                if parts:
                    lines.append(f"[{source} 表格 page{page}] " + "，".join(parts))

        return "\n".join(lines)


if __name__ == "__main__":
    loader = DocumentLoader()
    chunks = loader.load_directory(Path(__file__).parent.parent / "data" / "documents")
    print(f"\n总计: {len(chunks)} chunks")
    for c in chunks[:3]:
        print(f"  [{c.metadata['source']}] {c.content[:80]}...")
